"""Fill the interim report template with Module 1 (ABSA) content.

Only the sections that are this student's responsibility are filled.
Sections belonging to the other 3 team members (Preference Elicitation,
RAG/KG, Itinerary Planning) are left as template placeholders.

Output: interim_report_module1_filled.docx
"""

from docx import Document
from copy import deepcopy

SRC = "interim_report_template.docx"
DST = "interim_report_module1_filled_v2.docx"

# ---------------------------------------------------------------------------
# Module 1 content
# ---------------------------------------------------------------------------

CONTENT_6_2_DATA_PREP = (
    "The data preparation stage of the aspect extraction module has been fully implemented. "
    "A total of 48,364 user reviews covering 180 Sri Lankan Points of Interest (POIs) "
    "were collected from two complementary sources: the TripAdvisor public review pages, "
    "scraped via the Apify cloud-scraping service with a cap of 200 reviews per POI, "
    "and the publicly available Kaggle Sri Lanka tourism review corpus. "
    "The 180 POIs span all major Sri Lankan provinces and cover diverse categories "
    "including temples, national parks, beaches, waterfalls, and historical sites."
)

CONTENT_6_2_DATA_PREP_2 = (
    "The collected raw reviews were processed through a three-stage cleaning pipeline. "
    "First, text cleaning removed HTML tags, URLs, and emoji characters and normalized "
    "whitespace and punctuation, while preserving the original review text alongside the "
    "cleaned version. Second, language filtering using the langdetect library retained only "
    "English-language reviews with a confidence threshold of 0.7 and a minimum text length "
    "of 15 characters. Third, near-duplicate detection using TF-IDF vectorization with cosine "
    "similarity at threshold 0.95 removed copy-pasted or repeated submissions."
)

CONTENT_6_2_DATA_PREP_3 = (
    "Each cleaned review is persisted as a JSON object containing place name, original and "
    "cleaned text, rating, travel date, published date, and reviewer location. A flat CSV "
    "(_all_reviews.csv) and per-place JSON files form the input dataset for downstream aspect "
    "extraction and POI profile construction."
)

CONTENT_6_3_ASPECT = (
    "The aspect extraction module has been implemented as a three-step pipeline: "
    "(1) suggestion detection, (2) span extraction, and (3) normalization. "
    "The module focuses on three high-value travel aspects directly relevant to itinerary "
    "planning - best_time, crowd_level, and cost_level - although the framework is "
    "extensible to additional aspects with minimal code changes."
)

CONTENT_6_3_STEP1 = (
    "Step 1 - Suggestion Detection. A Sentence-BERT contrastive scoring approach is used to "
    "identify review sentences containing actionable advice for each aspect category. "
    "The pretrained model sentence-transformers/all-MiniLM-L6-v2 (22 MB, ~14k sentences/sec on "
    "CPU) encodes each review sentence into a 384-dimensional embedding. For each aspect "
    "category, a curated set of fact-dense positive anchor sentences is pre-encoded. A net "
    "contrastive score is computed as: score = max_similarity(sentence, positive_anchors) "
    "minus alpha times max_similarity(sentence, negative_anchors), where alpha = 0.5 and the "
    "negative anchors are generic travel praise sentences. The negative term is essential to "
    "suppress false positives caused by shared travel vocabulary. A sentence is tagged for an "
    "aspect when its net score exceeds the threshold of 0.38, selected through evaluation on a "
    "manually labelled validation set."
)

CONTENT_6_3_STEP2 = (
    "Step 2 - Span Extraction. For each tagged sentence, the actionable phrase is extracted "
    "using SpaCy en_core_web_sm with three complementary techniques: (a) a PhraseMatcher with "
    "curated phrase lists per aspect (for example, 'early morning', 'dry season', 'very crowded', "
    "'entry fee'); (b) Named-Entity Recognition for TIME and MONEY entities; and (c) dependency "
    "parsing to extract imperative-verb subtrees and avoid-objects (for example, 'avoid weekends'). "
    "A custom Matcher pattern is added for Sri Lankan Rupee amounts to capture phrases such as "
    "'Rs 1500', 'LKR 5000', and '200 rupees'."
)

CONTENT_6_3_STEP3 = (
    "Step 3 - Normalization. Extracted free-text spans are mapped to fixed categorical "
    "schemas using regular-expression pattern matching: TIME_OF_DAY (EARLY_MORNING / MID_MORNING / "
    "AFTERNOON / EVENING / NIGHT), SEASON (DRY_SEASON / MONSOON / SHOULDER), DAY_TYPE (WEEKEND / "
    "WEEKDAY / PUBLIC_HOLIDAY), CROWD_LEVEL (1-5 ordinal scale from EMPTY to PACKED with negation "
    "handling such as 'not crowded' which downgrades the level by two), and COST_LEVEL (FREE / "
    "LOW / MODERATE / HIGH / VERY_HIGH). For cost, two parallel signals are reported per review "
    "to preserve interpretability when they disagree: amount_level derived strictly from numeric "
    "LKR values, and sentiment_level derived strictly from opinion adjectives such as 'expensive' "
    "or 'cheap'."
)

CONTENT_6_3_AGGREGATION = (
    "POI Aggregation. The normalized values across all reviews of a place are aggregated using "
    "frequency-based majority voting. Each per-POI profile reports the dominant value, the "
    "number of supporting reviews, and a confidence score defined as the fraction of "
    "category-tagged reviews agreeing with the dominant value. The final output is 180 "
    "planner-ready JSON profiles, each containing best_time, crowd, and cost fields. "
    "These structured profiles serve as the input to the downstream Knowledge Graph, "
    "Preference Elicitation, and Itinerary Planning modules."
)

CONTENT_6_3_EVALUATION = (
    "Evaluation. The aspect classification step has been validated using a stratified gold-label "
    "evaluation set of 300 review sentences (150 system-tagged sentences for false-positive "
    "analysis and 150 system-untagged sentences for false-negative analysis), each independently "
    "labelled across all three aspect categories. The Sentence-BERT contrastive classifier "
    "achieves a macro F1-score of 0.756. Per-category F1 scores are: cost_level 0.933, "
    "best_time 0.785, and crowd_level 0.535. The cost_level result is strong because numeric "
    "amounts and explicit fee vocabulary provide unambiguous signals. The crowd_level F1 "
    "of 0.535 is identified as a known limitation - error analysis shows that "
    "general-purpose embeddings struggle to distinguish genuine crowd-density descriptions "
    "from sentences that merely share vocabulary such as 'tourist', 'group', and 'entrance'."
)

CONTENT_7_2_PROGRESS_M1 = (
    "Module 1 (Aspect-Based Sentiment Analysis) is fully implemented and operational on the "
    "180-POI dataset. The complete pipeline - data collection, preprocessing, suggestion "
    "detection, span extraction, normalization, and POI aggregation - executes end-to-end "
    "and produces planner-ready JSON profiles at output/aggregation/_poi_profiles.json. "
    "A gold-label evaluation framework with 300 manually labelled sentences has been built "
    "and used to validate aspect classification accuracy (macro F1 = 0.756). "
    "Modules 2-4 (Preference Elicitation, RAG with Knowledge Graph, and Itinerary Planning) "
    "are progressing in parallel by other team members."
)

# 2.3 Aspect-Based Sentiment Analysis in Tourism (literature review)
CONTENT_2_3_ABSA_LIT = (
    "Aspect-Based Sentiment Analysis (ABSA) is a fine-grained sentiment analysis task that "
    "extracts opinions about specific aspects of an entity rather than treating the review as "
    "a single sentiment unit. In the tourism domain, a single review often expresses multiple, "
    "potentially conflicting opinions in one sentence, such as describing a beach as having a "
    "beautiful sunset but also being overcrowded in the evening. Coarse-grained sentiment "
    "classifiers cannot capture this nuance, motivating the use of ABSA for travel reviews."
)

CONTENT_2_3_ABSA_LIT_2 = (
    "Recent studies on tourism review mining have applied ABSA to extract aspects such as "
    "scenery, crowd density, cleanliness, safety, accessibility, family suitability, activities, "
    "and cost. Approaches range from supervised neural sequence taggers (BiLSTM-CRF, BERT-based "
    "token classifiers) to unsupervised topic-modelling and embedding-based methods. Supervised "
    "methods generally produce higher accuracy but require labelled data which is scarce for "
    "Sri Lankan tourism. Embedding-based zero-shot methods - particularly Sentence-BERT "
    "(Reimers and Gurevych, 2019) with contrastive scoring against curated anchor sentences - "
    "have been shown to perform competitively without labelled training data and are therefore "
    "well-suited to the data-scarce setting of this project."
)

CONTENT_2_3_ABSA_LIT_3 = (
    "A consistent finding in the literature is that ABSA outputs must be normalized to "
    "structured, comparable values to be useful for downstream recommendation tasks. Free-text "
    "phrases such as 'go early', 'before 9 am', and 'sunrise' all describe the same actionable "
    "concept and should map to a single label such as EARLY_MORNING. This project follows "
    "that principle by combining sentence-level aspect classification with span extraction "
    "and rule-based normalization to produce planner-ready POI profiles."
)

# 3.2 - additional paragraph specific to our adoption
CONTENT_3_2_EXTRA = (
    "In this project, ABSA is implemented using a Sentence-BERT contrastive scoring approach "
    "for aspect classification, followed by SpaCy-based span extraction and rule-based "
    "normalization to fixed schemas. This combination is chosen because it requires no "
    "labelled training data, runs efficiently on CPU, and produces structured outputs that "
    "can be directly consumed by the downstream Knowledge Graph, RAG, and Itinerary Planning "
    "modules. The approach was validated on 300 manually labelled sentences, achieving a "
    "macro F1 of 0.756."
)

# 5.6.1 actual review chunk structure used in this project
CONTENT_5_6_1_STRUCTURE = (
    '{\n'
    '  "place": "Sigiriya_Lion_Rock",\n'
    '  "text_clean": "Go early in the morning to avoid the heat and crowds.",\n'
    '  "analysis_tags": {"best_time": true, "crowd_level": true, "cost_level": false},\n'
    '  "analysis_scores": {"best_time": 0.52, "crowd_level": 0.41, "cost_level": 0.12},\n'
    '  "analysis_sentences": {\n'
    '    "best_time": ["Go early in the morning to avoid the heat and crowds."],\n'
    '    "crowd_level": ["Go early in the morning to avoid the heat and crowds."]\n'
    '  },\n'
    '  "extracted_spans": {\n'
    '    "best_time": {"time_spans": ["early in the morning"], "action": "go", "avoid": ["the heat and crowds"]},\n'
    '    "crowd_level": {"crowd_spans": ["crowds"], "crowd_when": ["in the morning"]}\n'
    '  },\n'
    '  "normalized": {\n'
    '    "best_time": {"time_of_day": "EARLY_MORNING", "season": null, "avoid": []},\n'
    '    "crowd_level": {"crowd_level": 4, "crowd_label": "BUSY", "crowd_when": []}\n'
    '  }\n'
    '}'
)

# 7.5 Challenges (Module 1 specific)
CONTENT_7_5_CHALLENGES = (
    "Several challenges were encountered during Module 1 implementation. (1) Noisy review text: "
    "user reviews contain spelling mistakes, mixed languages (Sinhala-English code switching), "
    "and informal abbreviations that affect sentence-level classification. Language filtering "
    "and TF-IDF deduplication mitigated but did not eliminate these issues. "
    "(2) Vocabulary overlap across aspects: words such as 'tourist', 'entrance', and 'group' "
    "appear in multiple aspect contexts, producing false positives in the crowd_level classifier "
    "(F1 = 0.535). Attempts to add category-specific negative anchors improved precision but "
    "reduced recall, yielding a lower overall F1. This indicates that general-purpose embeddings "
    "have a fundamental limit on this discrimination task and motivates future fine-tuning. "
    "(3) Threshold selection: choosing the classification threshold required a labelled "
    "validation set; this was addressed by building a stratified gold-label evaluation set "
    "and selecting threshold 0.38 by F1 maximization. "
    "(4) Cost signal disagreement: amount-based and sentiment-based cost signals can conflict "
    "(for example, a review saying the Rs 100 fee is 'very expensive'). The current implementation "
    "reports both signals separately to preserve interpretability. "
    "(5) Evaluation dataset preparation: manually labelling 300 sentences across three categories "
    "is time-consuming; a custom Streamlit labelling interface was built to accelerate this work."
)

# Appendix A content for the user's individual contribution
INDIVIDUAL_CONTRIBUTION = [
    "Designed and implemented Module 1 - Aspect-Based Sentiment Analysis - which provides "
    "the foundational structured POI profiles consumed by the other three modules.",

    "Collected and preprocessed a multi-source review corpus of 48,364 English reviews "
    "covering 180 Sri Lankan POIs (TripAdvisor via Apify and Kaggle).",

    "Implemented sentence-level suggestion classification using a Sentence-BERT "
    "(all-MiniLM-L6-v2) contrastive scoring approach with curated positive and negative "
    "anchor sets and an evaluation-justified threshold of 0.38.",

    "Implemented span extraction using SpaCy (PhraseMatcher, Named-Entity Recognition, "
    "and dependency parsing) and rule-based normalization to fixed categorical schemas "
    "for time-of-day, season, day-type, crowd level (1-5), and cost level.",

    "Implemented POI aggregation with frequency-based majority voting and confidence "
    "scoring, producing 180 planner-ready JSON profiles.",

    "Built a gold-label evaluation framework consisting of (a) a stratified sentence "
    "sampler, (b) a Streamlit-based labelling interface with category-specific guidance, "
    "and (c) a precision/recall/F1 scorer with error analysis. Manually labelled 300 "
    "sentences across the three aspect categories.",

    "Reported macro F1 of 0.756 on the gold-labelled evaluation set "
    "(cost_level F1 = 0.933, best_time F1 = 0.785, crowd_level F1 = 0.535) and "
    "performed an error analysis identifying the crowd_level vocabulary-overlap issue "
    "as a direction for future work.",
]

# ---------------------------------------------------------------------------
# Replacement helpers
# ---------------------------------------------------------------------------

def replace_paragraph_text(para, new_text: str):
    """Replace text in a paragraph while preserving the first run's formatting."""
    if not para.runs:
        para.add_run(new_text)
        return
    # Keep first run, clear others
    first = para.runs[0]
    first.text = new_text
    for run in para.runs[1:]:
        run.text = ""


def insert_paragraph_after(para, text, style=None):
    """Insert a new paragraph with `text` directly after `para`."""
    new_p = deepcopy(para)
    # Clear runs in the copy
    for run in new_p.runs:
        run.text = ""
    if new_p.runs:
        new_p.runs[0].text = text
    else:
        new_p.add_run(text)
    if style is not None:
        new_p.style = style
    para._element.addnext(new_p._element)
    return new_p


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document(SRC)

    # Build a list of (index, paragraph) for easier scanning
    paragraphs = list(doc.paragraphs)

    # ------- 6.2 Data Preparation -------
    placeholder_62 = ("[Describe collected review data, cleaning, JSON/CSV structure, "
                      "POI identifiers, chunking, and aspect list preparation.]")
    for i, p in enumerate(paragraphs):
        if placeholder_62 in p.text:
            replace_paragraph_text(p, CONTENT_6_2_DATA_PREP)
            p2 = insert_paragraph_after(p, CONTENT_6_2_DATA_PREP_2)
            insert_paragraph_after(p2, CONTENT_6_2_DATA_PREP_3)
            print(f"Filled section 6.2 at paragraph {i}")
            break

    # Re-fetch paragraphs because we just inserted new ones
    paragraphs = list(doc.paragraphs)

    # ------- 2.3 Aspect-Based Sentiment Analysis (literature review) -------
    placeholder_23 = ("[Explain how ABSA extracts fine-grained opinions from reviews. "
                      "Discuss aspects such as scenery, crowd, cleanliness, safety, "
                      "accessibility, family suitability, activities, and cost. "
                      "Cite ABSA and tourism review mining studies.]")
    for i, p in enumerate(paragraphs):
        if placeholder_23 in p.text:
            replace_paragraph_text(p, CONTENT_2_3_ABSA_LIT)
            p2 = insert_paragraph_after(p, CONTENT_2_3_ABSA_LIT_2)
            insert_paragraph_after(p2, CONTENT_2_3_ABSA_LIT_3)
            print(f"Filled section 2.3 at paragraph {i}")
            break

    paragraphs = list(doc.paragraphs)

    # ------- 3.2 Aspect-Based Sentiment Analysis (append our specifics) -------
    marker_32 = ("Aspect-based sentiment analysis is adapted because travel reviews "
                 "usually contain several opinions within a single sentence")
    for i, p in enumerate(paragraphs):
        if marker_32 in p.text:
            insert_paragraph_after(p, CONTENT_3_2_EXTRA)
            print(f"Appended Module 1 specifics to section 3.2 at paragraph {i}")
            break

    paragraphs = list(doc.paragraphs)

    # ------- 5.6.1 Review Chunk Structure (replace generic example) -------
    marker_561 = '"chunk_id": "C001"'
    for i, p in enumerate(paragraphs):
        if marker_561 in p.text:
            # Replace the entire JSON block with our actual structure
            replace_paragraph_text(p, CONTENT_5_6_1_STRUCTURE)
            print(f"Filled section 5.6.1 at paragraph {i}")
            break

    paragraphs = list(doc.paragraphs)

    # ------- 7.5 Challenges -------
    placeholder_75 = ("[Discuss expected or encountered challenges such as noisy reviews, "
                      "subjective statements, missing metadata, LLM cost, API rate limits, "
                      "evidence quality, and evaluation dataset preparation.]")
    for i, p in enumerate(paragraphs):
        if placeholder_75 in p.text:
            replace_paragraph_text(p, CONTENT_7_5_CHALLENGES)
            print(f"Filled section 7.5 at paragraph {i}")
            break

    paragraphs = list(doc.paragraphs)

    # ------- 6.3 Aspect Extraction -------
    placeholder_63 = ("[Describe the current ABSA implementation or prototype. "
                      "Mention whether LLM API, Hugging Face model, manual labels, "
                      "or prompt-based extraction is used.]")
    for i, p in enumerate(paragraphs):
        if placeholder_63 in p.text:
            replace_paragraph_text(p, CONTENT_6_3_ASPECT)
            p1 = insert_paragraph_after(p,  CONTENT_6_3_STEP1)
            p2 = insert_paragraph_after(p1, CONTENT_6_3_STEP2)
            p3 = insert_paragraph_after(p2, CONTENT_6_3_STEP3)
            p4 = insert_paragraph_after(p3, CONTENT_6_3_AGGREGATION)
            insert_paragraph_after(p4, CONTENT_6_3_EVALUATION)
            print(f"Filled section 6.3 at paragraph {i}")
            break

    # Re-fetch
    paragraphs = list(doc.paragraphs)

    # ------- 7.2 Current Progress (append our module status) -------
    marker_72 = "At the interim stage, the project has completed problem identification"
    for i, p in enumerate(paragraphs):
        if marker_72 in p.text:
            insert_paragraph_after(p, CONTENT_7_2_PROGRESS_M1)
            print(f"Appended Module 1 progress note after paragraph {i}")
            break

    # Re-fetch
    paragraphs = list(doc.paragraphs)

    # ------- Appendix A: Individual Contribution -------
    # Find "Contribution:" line
    for i, p in enumerate(paragraphs):
        if p.text.strip() == "Contribution:":
            # Insert each bullet point
            anchor = p
            for line in INDIVIDUAL_CONTRIBUTION:
                anchor = insert_paragraph_after(anchor, "- " + line)
            print(f"Filled Appendix A contribution after paragraph {i}")
            break

    # ------- Update student name in front-matter tables (Table 0 and Table 1) -------
    student_name  = "Bhathiya Ranasinghe"
    student_index = "[Index Number]"  # leave as placeholder
    # Only fill the FIRST row of Table 0 and Table 1 (top of group list)
    for table_idx in [0, 1]:
        if table_idx < len(doc.tables):
            row0 = doc.tables[table_idx].rows[0]
            row0.cells[0].text = student_index
            row0.cells[1].text = student_name
            print(f"Set student name in Table {table_idx}, row 0")

    # ------- Update Module 1 row in Table 5 (Section 1.5) for clarity -------
    # Make the first row reflect Module 1 specifics
    if len(doc.tables) > 5:
        t5 = doc.tables[5]
        if len(t5.rows) >= 2:
            t5.rows[1].cells[0].text = "Aspect-Based Sentiment Analysis (Module 1)"
            t5.rows[1].cells[1].text = (
                "Extracts structured travel aspects (best_time, crowd_level, cost_level) "
                "from user-generated reviews and aggregates them into per-POI profiles "
                "with confidence scores."
            )
            print("Updated Table 5 row 1 (Module 1)")

    # ------- Update Table 10 (5.3 Module Descriptions) - rows for Module 1 -------
    if len(doc.tables) > 10:
        t10 = doc.tables[10]
        # Row 1: Data Preprocessing
        if len(t10.rows) >= 2:
            t10.rows[1].cells[1].text = ("Cleans raw reviews; applies HTML/URL/emoji removal, "
                                          "language filtering, and TF-IDF deduplication.")
            t10.rows[1].cells[2].text = "Raw TripAdvisor and Kaggle reviews (48,364 reviews, 180 POIs)"
            t10.rows[1].cells[3].text = "Cleaned per-place JSON files and combined CSV"
        # Row 2: ABSA Module
        if len(t10.rows) >= 3:
            t10.rows[2].cells[1].text = ("Detects, extracts, normalizes, and aggregates aspects "
                                          "(best_time, crowd_level, cost_level) using SBERT + SpaCy.")
            t10.rows[2].cells[2].text = "Cleaned reviews"
            t10.rows[2].cells[3].text = "Per-POI structured profiles with confidence scores"
        print("Updated Table 10 rows 1 and 2 (Module 1)")

    # ------- Update Table 13 (6.7 Tools) -------
    if len(doc.tables) > 13:
        t13 = doc.tables[13]
        # Replace generic NLP/ABSA row
        for row in t13.rows[1:]:
            if row.cells[0].text.strip().lower().startswith("nlp"):
                row.cells[1].text = ("sentence-transformers (all-MiniLM-L6-v2) for "
                                      "sentence embeddings; SpaCy en_core_web_sm for NER, "
                                      "PhraseMatcher, and dependency parsing")
                print("Updated Table 13 NLP/ABSA row")
                break

    # ------- Update Table 14 (7.4 Evaluation Plan) - ABSA row with actual results -------
    if len(doc.tables) > 14:
        t14 = doc.tables[14]
        # Row 1 should be ABSA
        if len(t14.rows) >= 2:
            t14.rows[1].cells[1].text = ("Compare extracted aspect tags with manually labelled "
                                          "ground-truth on a 300-sentence stratified gold set "
                                          "(150 system-tagged, 150 untagged)")
            t14.rows[1].cells[2].text = ("Precision, Recall, F1-score per aspect; macro F1. "
                                          "Current results: macro F1 = 0.756 "
                                          "(cost_level 0.933, best_time 0.785, crowd_level 0.535)")
            print("Updated Table 14 row 1 (ABSA evaluation)")

    # ------- Update Appendix A name field -------
    for i, p in enumerate(paragraphs):
        if p.text.strip().startswith("Name of student:"):
            replace_paragraph_text(p, f"Name of student: {student_name}")
            print(f"Set student name in Appendix A")
            break

    doc.save(DST)
    print(f"\nSaved: {DST}")


if __name__ == "__main__":
    main()
